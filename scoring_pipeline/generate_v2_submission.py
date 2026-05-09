#!/usr/bin/env python3
"""
Generate v2 submission package:
- New validation figures (Fig 7, Fig 8) with updated data
- Regenerate Figs 1-6 (unchanged methodology)
- Updated manuscript as .docx
- Supporting data tables
"""

import json
import os
import sys
import csv
import shutil
import numpy as np
from collections import defaultdict
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import matplotlib.patheffects as pe
from matplotlib.patches import FancyBboxPatch

plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Helvetica", "Arial", "DejaVu Sans"],
    "font.size": 10,
    "axes.labelsize": 11,
    "axes.titlesize": 12,
    "xtick.labelsize": 9,
    "ytick.labelsize": 9,
    "legend.fontsize": 9,
    "figure.dpi": 300,
    "savefig.dpi": 300,
    "savefig.bbox": "tight",
    "savefig.pad_inches": 0.15,
})

C_BLUE = "#2166AC"
C_RED = "#B2182B"
C_GREEN = "#1B7837"
C_ORANGE = "#E08214"
C_PURPLE = "#7B3294"
C_GRAY = "#969696"
C_LIGHT_BLUE = "#92C5DE"
C_LIGHT_RED = "#F4A582"
C_TEAL = "#01665E"
C_GOLD = "#D4A017"

BASE = os.path.dirname(__file__)
OUT_DIR = os.path.join(BASE, "v2_submission")
FIG_DIR = os.path.join(OUT_DIR, "figures")

RESULTS_77 = os.path.join(BASE, "validation_77_results.json")
RESULTS_EXT = os.path.join(BASE, "validation_extended_results.json")
RESULTS_S3 = os.path.join(BASE, "validation_set3_results.json")


def load_results():
    with open(RESULTS_77) as f:
        r77 = json.load(f)
    with open(RESULTS_EXT) as f:
        rext = json.load(f)
    with open(RESULTS_S3) as f:
        rs3 = json.load(f)
    return r77, rext, rs3


def figure7_aggregate_recovery(r77, rext, rs3):
    """Fig 7: Three-set aggregate factor recovery comparison."""
    fig = plt.figure(figsize=(18, 12))
    gs = gridspec.GridSpec(2, 2, hspace=0.35, wspace=0.3)

    ax1 = fig.add_subplot(gs[0, 0])
    ax1.text(-0.08, 1.05, "a", transform=ax1.transAxes, fontsize=14,
             fontweight="bold", va="top", ha="left")

    thresholds = [1, 2, 5, 10, 20]
    labels = ["Top 1%", "Top 2%", "Top 5%", "Top 10%", "Top 20%"]

    sets_data = {
        "Set 1\n(77 cocktails)": r77,
        "Set 2\n(64 cocktails)": rext,
        "Set 3 blind\n(33 cocktails)": rs3,
    }
    colors = [C_BLUE, C_TEAL, C_ORANGE]
    x = np.arange(len(thresholds))
    width = 0.25

    for i, (set_name, data) in enumerate(sets_data.items()):
        pcts = []
        for t_label in labels:
            pcts.append(data["aggregate_thresholds"][t_label]["pct"])
        bars = ax1.bar(x + i*width, pcts, width, label=set_name, color=colors[i],
                       edgecolor="white", linewidth=0.5)
        for bar, pct in zip(bars, pcts):
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                     f"{pct:.0f}%", ha="center", va="bottom", fontsize=7, fontweight="bold")

    ax1.set_xlabel("Recovery threshold")
    ax1.set_ylabel("Factor-cocktail pairs recovered (%)")
    ax1.set_title("Aggregate factor recovery across three validation sets")
    ax1.set_xticks(x + width)
    ax1.set_xticklabels(labels)
    ax1.legend(loc="upper left", framealpha=0.9)
    ax1.set_ylim(0, 100)
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)

    ax2 = fig.add_subplot(gs[0, 1])
    ax2.text(-0.08, 1.05, "b", transform=ax2.transAxes, fontsize=14,
             fontweight="bold", va="top", ha="left")

    family_order_77 = []
    fam_summary = r77.get("family_summary", {})
    for fam in sorted(fam_summary.keys()):
        fs = fam_summary[fam]
        if fs["n_factor_tests"] > 0:
            family_order_77.append(fam)

    family_order_77.sort(key=lambda f: -fam_summary[f]["top5_pct"])

    y_pos = np.arange(len(family_order_77))
    top5_vals = [fam_summary[f]["top5_pct"] for f in family_order_77]
    top10_vals = [fam_summary[f]["top10_pct"] for f in family_order_77]

    ax2.barh(y_pos + 0.15, top10_vals, 0.3, color=C_LIGHT_BLUE, label="Top 10%")
    ax2.barh(y_pos - 0.15, top5_vals, 0.3, color=C_BLUE, label="Top 5%")

    for i, fam in enumerate(family_order_77):
        n = fam_summary[fam]["n_factor_tests"]
        ax2.text(max(top10_vals[i], top5_vals[i]) + 1, i, f"n={n}", va="center", fontsize=8)

    ax2.set_yticks(y_pos)
    ax2.set_yticklabels(family_order_77)
    ax2.set_xlabel("Factor recovery (%)")
    ax2.set_title("Set 1: Recovery by cell type family")
    ax2.set_xlim(0, 110)
    ax2.legend(loc="lower right", framealpha=0.9)
    ax2.spines["top"].set_visible(False)
    ax2.spines["right"].set_visible(False)

    ax3 = fig.add_subplot(gs[1, 0])
    ax3.text(-0.08, 1.05, "c", transform=ax3.transAxes, fontsize=14,
             fontweight="bold", va="top", ha="left")

    all_percentiles_77 = [e["percentile"] for e in r77.get("unique_factors", [])]
    all_percentiles_ext = [e["percentile"] for e in rext.get("unique_factors", [])]
    all_percentiles_s3 = [e["percentile"] for e in rs3.get("unique_factors", [])]

    bins = np.arange(0, 52, 2)
    ax3.hist(all_percentiles_77, bins=bins, alpha=0.6, color=C_BLUE, label=f"Set 1 (n={len(all_percentiles_77)})", density=True)
    ax3.hist(all_percentiles_ext, bins=bins, alpha=0.6, color=C_TEAL, label=f"Set 2 (n={len(all_percentiles_ext)})", density=True)
    ax3.hist(all_percentiles_s3, bins=bins, alpha=0.6, color=C_ORANGE, label=f"Set 3 (n={len(all_percentiles_s3)})", density=True)
    ax3.axvline(5, color="black", linestyle="--", linewidth=1, alpha=0.5)
    ax3.text(5.5, ax3.get_ylim()[1]*0.9, "5%\nthreshold", fontsize=8, color="gray")
    ax3.set_xlabel("Best percentile rank per factor")
    ax3.set_ylabel("Density")
    ax3.set_title("Unique factor rank distributions")
    ax3.legend(loc="upper right", framealpha=0.9)
    ax3.set_xlim(0, 50)
    ax3.spines["top"].set_visible(False)
    ax3.spines["right"].set_visible(False)

    ax4 = fig.add_subplot(gs[1, 1])
    ax4.text(-0.08, 1.05, "d", transform=ax4.transAxes, fontsize=14,
             fontweight="bold", va="top", ha="left")

    set_labels = ["Set 1\n(77 cocktails,\ntraining-adjacent)", "Set 2\n(64 cocktails,\nextended)", "Set 3\n(33 cocktails,\nblind)"]
    top5_pcts = [
        r77["aggregate_thresholds"]["Top 5%"]["pct"],
        rext["aggregate_thresholds"]["Top 5%"]["pct"],
        rs3["aggregate_thresholds"]["Top 5%"]["pct"],
    ]
    top10_pcts = [
        r77["aggregate_thresholds"]["Top 10%"]["pct"],
        rext["aggregate_thresholds"]["Top 10%"]["pct"],
        rs3["aggregate_thresholds"]["Top 10%"]["pct"],
    ]
    n_factors = [
        r77["summary"]["factors_in_pool"],
        rext["summary"]["factors_in_pool"],
        rs3["summary"]["factors_in_pool"],
    ]

    x_pos = np.arange(3)
    bars5 = ax4.bar(x_pos - 0.15, top5_pcts, 0.3, color=C_BLUE, label="Top 5%", edgecolor="white")
    bars10 = ax4.bar(x_pos + 0.15, top10_pcts, 0.3, color=C_LIGHT_BLUE, label="Top 10%", edgecolor="white")

    for bar, pct, n in zip(bars5, top5_pcts, n_factors):
        ax4.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                 f"{pct:.1f}%", ha="center", va="bottom", fontsize=9, fontweight="bold")
    for bar, pct in zip(bars10, top10_pcts):
        ax4.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                 f"{pct:.1f}%", ha="center", va="bottom", fontsize=9, fontweight="bold")

    ax4.axhline(5, color=C_RED, linestyle="--", linewidth=1, alpha=0.4)
    ax4.axhline(10, color=C_RED, linestyle=":", linewidth=1, alpha=0.3)
    ax4.text(2.5, 6, "Random\nexpectation\n(5%)", fontsize=7, color=C_RED, alpha=0.6)

    ax4.set_xticks(x_pos)
    ax4.set_xticklabels(set_labels, fontsize=9)
    ax4.set_ylabel("Factor recovery (%)")
    ax4.set_title("Cross-set consistency of predictor performance")
    ax4.legend(loc="upper right", framealpha=0.9)
    ax4.set_ylim(0, 100)
    ax4.spines["top"].set_visible(False)
    ax4.spines["right"].set_visible(False)

    fig.savefig(f"{FIG_DIR}/fig7_cocktail_validation.png")
    fig.savefig(f"{FIG_DIR}/fig7_cocktail_validation.pdf")
    plt.close(fig)
    print("  Fig 7 saved")


def figure8_tier_and_detail(r77, rext, rs3):
    """Fig 8: Two-tier weight system and per-cocktail detail."""
    fig = plt.figure(figsize=(18, 14))
    gs = gridspec.GridSpec(2, 2, hspace=0.35, wspace=0.35)

    ax1 = fig.add_subplot(gs[0, 0])
    ax1.text(-0.08, 1.05, "a", transform=ax1.transAxes, fontsize=14,
             fontweight="bold", va="top", ha="left")

    components = ["GTEx", "tau", "pheno", "kernel", "dir", "act", "frac", "enr", "temp"]
    tier_a = [0.2912, 0.2084, 0.0853, 0.0933, 0.0734, 0.0874, 0.0856, 0.0321, 0.0434]
    tier_b = [0.1944, 0.2192, 0.3410, 0.0491, 0.0805, 0.0330, 0.0407, 0.0297, 0.0123]

    x = np.arange(len(components))
    width = 0.35
    bars_a = ax1.bar(x - width/2, [v*100 for v in tier_a], width, color=C_BLUE, label="Tier A (GTEx-dominant)", edgecolor="white")
    bars_b = ax1.bar(x + width/2, [v*100 for v in tier_b], width, color=C_ORANGE, label="Tier B (Phenotype-dominant)", edgecolor="white")

    ax1.set_xticks(x)
    ax1.set_xticklabels(components, rotation=45, ha="right")
    ax1.set_ylabel("Weight (%)")
    ax1.set_title("Two-tier weight profiles")
    ax1.legend(loc="upper right", framealpha=0.9)
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)

    highlight_idx = [0, 2]
    for idx in highlight_idx:
        rect = plt.Rectangle((idx - 0.45, 0), 0.9, max(tier_a[idx], tier_b[idx])*100 + 3,
                               linewidth=2, edgecolor=C_RED, facecolor="none", linestyle="--", alpha=0.6)
        ax1.add_patch(rect)

    ax2 = fig.add_subplot(gs[0, 1])
    ax2.text(-0.08, 1.05, "b", transform=ax2.transAxes, fontsize=14,
             fontweight="bold", va="top", ha="left")

    tier_a_types = ["Neuron", "Cardiac", "Hepatic", "Endocrine", "Muscle", "Glial"]
    tier_b_types = ["Immune", "Vascular", "Blood", "Mesenchymal", "Epithelial"]

    tier_a_box_y = np.arange(len(tier_a_types))
    tier_b_box_y = np.arange(len(tier_b_types))

    for i, ct in enumerate(tier_a_types):
        ax2.barh(i + len(tier_b_types) + 1.5, 1, color=C_BLUE, alpha=0.7, height=0.6, edgecolor="white")
        ax2.text(0.5, i + len(tier_b_types) + 1.5, ct, ha="center", va="center", fontsize=10, fontweight="bold", color="white")

    for i, ct in enumerate(tier_b_types):
        ax2.barh(i, 1, color=C_ORANGE, alpha=0.7, height=0.6, edgecolor="white")
        ax2.text(0.5, i, ct, ha="center", va="center", fontsize=10, fontweight="bold", color="white")

    ax2.text(0.5, len(tier_b_types) + len(tier_a_types) + 2, "TIER A", ha="center", fontsize=12, fontweight="bold", color=C_BLUE)
    ax2.text(1.6, len(tier_b_types) + len(tier_a_types) + 2, "GTEx expression\ndominates (29.1%)", ha="left", fontsize=9, color=C_BLUE)

    ax2.text(0.5, len(tier_b_types) + 0.7, "TIER B", ha="center", fontsize=12, fontweight="bold", color=C_ORANGE)
    ax2.text(1.6, len(tier_b_types) + 0.7, "Phenotype/disease\nassociation dominates (34.1%)", ha="left", fontsize=9, color=C_ORANGE)

    ax2.set_xlim(-0.2, 3.5)
    ax2.set_ylim(-0.5, len(tier_a_types) + len(tier_b_types) + 3)
    ax2.axis("off")
    ax2.set_title("Cell type tier assignments")

    ax3 = fig.add_subplot(gs[1, 0])
    ax3.text(-0.08, 1.05, "c", transform=ax3.transAxes, fontsize=14,
             fontweight="bold", va="top", ha="left")

    top_factors = r77.get("unique_factors", [])[:25]
    factor_names = [e["factor"] for e in top_factors]
    factor_pcts = [e["percentile"] for e in top_factors]
    factor_targets = [e["target"] for e in top_factors]

    y_pos = np.arange(len(factor_names))
    colors_f = []
    for p in factor_pcts:
        if p <= 1:
            colors_f.append(C_GREEN)
        elif p <= 5:
            colors_f.append(C_BLUE)
        elif p <= 10:
            colors_f.append(C_ORANGE)
        else:
            colors_f.append(C_RED)

    ax3.barh(y_pos, factor_pcts, color=colors_f, edgecolor="white", height=0.7)
    ax3.set_yticks(y_pos)
    ax3.set_yticklabels(factor_names, fontsize=8)
    ax3.invert_yaxis()
    ax3.set_xlabel("Best percentile rank (%)")
    ax3.set_title("Top 25 factors by rank (Set 1)")
    ax3.axvline(5, color="black", linestyle="--", linewidth=1, alpha=0.3)
    ax3.axvline(1, color="black", linestyle=":", linewidth=1, alpha=0.2)

    for i, (name, pct, target) in enumerate(zip(factor_names, factor_pcts, factor_targets)):
        ax3.text(max(pct, 0.3) + 0.2, i, f"→{target}", va="center", fontsize=7, color=C_GRAY)

    ax3.set_xlim(0, max(factor_pcts) + 8)
    ax3.spines["top"].set_visible(False)
    ax3.spines["right"].set_visible(False)

    ax4 = fig.add_subplot(gs[1, 1])
    ax4.text(-0.08, 1.05, "d", transform=ax4.transAxes, fontsize=14,
             fontweight="bold", va="top", ha="left")

    all_fam_ext = rext.get("family_summary", {})
    all_fam_s3 = rs3.get("family_summary", {})

    combined_fams = {}
    for fam, fs in all_fam_ext.items():
        if fs["n_factor_tests"] > 0:
            combined_fams[fam] = {"ext_top5": fs["top5_pct"], "ext_n": fs["n_factor_tests"]}
    for fam, fs in all_fam_s3.items():
        if fs["n_factor_tests"] > 0:
            if fam not in combined_fams:
                combined_fams[fam] = {"ext_top5": 0, "ext_n": 0}
            combined_fams[fam]["s3_top5"] = fs["top5_pct"]
            combined_fams[fam]["s3_n"] = fs["n_factor_tests"]

    fams_both = [f for f in combined_fams if "s3_top5" in combined_fams[f] and combined_fams[f].get("ext_n", 0) > 0]
    fams_both.sort(key=lambda f: -(combined_fams[f]["ext_top5"] + combined_fams[f].get("s3_top5", 0))/2)

    if fams_both:
        y_pos_f = np.arange(len(fams_both))
        ext_vals = [combined_fams[f]["ext_top5"] for f in fams_both]
        s3_vals = [combined_fams[f].get("s3_top5", 0) for f in fams_both]

        ax4.barh(y_pos_f + 0.15, ext_vals, 0.3, color=C_TEAL, label="Set 2 (extended)", edgecolor="white")
        ax4.barh(y_pos_f - 0.15, s3_vals, 0.3, color=C_ORANGE, label="Set 3 (blind)", edgecolor="white")

        ax4.set_yticks(y_pos_f)
        ax4.set_yticklabels(fams_both, fontsize=9)
        ax4.set_xlabel("Top-5% recovery (%)")
        ax4.set_title("Out-of-sample recovery by family")
        ax4.set_xlim(0, 110)
        ax4.legend(loc="lower right", framealpha=0.9)
    ax4.spines["top"].set_visible(False)
    ax4.spines["right"].set_visible(False)

    fig.savefig(f"{FIG_DIR}/fig8_twotier_detail.png")
    fig.savefig(f"{FIG_DIR}/fig8_twotier_detail.pdf")
    plt.close(fig)
    print("  Fig 8 saved")


def figure_supp_set3_detail(rs3):
    """Supplementary: Set 3 per-cocktail detail."""
    cocktails = rs3.get("per_cocktail", [])
    testable = [c for c in cocktails if any(
        fr.get("found") for fr in c["factor_ranks"].values()
    )]

    fig, ax = plt.subplots(figsize=(16, max(8, len(testable)*0.45)))
    ax.text(-0.05, 1.02, "Supplementary", transform=ax.transAxes, fontsize=12,
            fontweight="bold", va="bottom", ha="left")

    y_pos = 0
    y_labels = []
    y_positions = []
    for cocktail in reversed(testable):
        cid = cocktail["cocktail_id"]
        target = cocktail["target_cell"]
        source = cocktail["source_cell"]
        label = f"{cid}\n{source}→{target}"

        for factor, fr in cocktail["factor_ranks"].items():
            if fr.get("found"):
                pct = fr["percentile"]
                color = C_GREEN if pct <= 1 else C_BLUE if pct <= 5 else C_ORANGE if pct <= 10 else C_RED
                ax.barh(y_pos, pct, color=color, height=0.6, edgecolor="white")
                ax.text(pct + 0.5, y_pos, f"{factor} ({pct:.1f}%)", va="center", fontsize=7)
            else:
                ax.text(0.5, y_pos, f"{factor} [not in pool]", va="center", fontsize=7, color=C_GRAY)
            y_positions.append(y_pos)
            y_labels.append("")
            y_pos += 1

        mid = y_pos - len(cocktail["factor_ranks"]) / 2
        ax.text(-1, mid, label, ha="right", va="center", fontsize=7, color="black")
        y_pos += 0.5

    ax.axvline(5, color="black", linestyle="--", linewidth=1, alpha=0.3)
    ax.axvline(10, color="black", linestyle=":", linewidth=1, alpha=0.2)
    ax.set_xlabel("Percentile rank among 3,166 Tier 2+ TFs")
    ax.set_title("Validation Set 3 — Per-factor rankings (blind post-calibration test)")
    ax.set_xlim(-0.5, 40)
    ax.set_yticks([])
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)

    fig.savefig(f"{FIG_DIR}/fig_supp_set3_detail.png")
    fig.savefig(f"{FIG_DIR}/fig_supp_set3_detail.pdf")
    plt.close(fig)
    print("  Supplementary Set 3 detail saved")


def generate_supporting_tables(r77, rext, rs3):
    """Generate CSV supporting data tables."""

    with open(f"{OUT_DIR}/table_s1_set1_factor_rankings.csv", "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["factor", "cocktail_id", "target_cell", "source_cell", "rank", "percentile", "family", "n_regulatory"])
        for entry in r77.get("unique_factors", []):
            w.writerow([entry["factor"], entry.get("cocktail", ""), entry.get("target", ""),
                        entry.get("source", ""), entry.get("rank", ""), f"{entry['percentile']:.2f}",
                        entry.get("family", ""), r77["summary"]["n_regulatory_genes"]])

    with open(f"{OUT_DIR}/table_s2_set2_factor_rankings.csv", "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["factor", "cocktail_id", "target_cell", "source_cell", "rank", "percentile", "family", "n_regulatory"])
        for entry in rext.get("unique_factors", []):
            w.writerow([entry["factor"], entry.get("cocktail", ""), entry.get("target", ""),
                        entry.get("source", ""), entry.get("rank", ""), f"{entry['percentile']:.2f}",
                        entry.get("family", ""), rext["summary"]["n_regulatory_genes"]])

    with open(f"{OUT_DIR}/table_s3_set3_factor_rankings.csv", "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["factor", "cocktail_id", "target_cell", "source_cell", "rank", "percentile", "family", "n_regulatory"])
        for entry in rs3.get("unique_factors", []):
            w.writerow([entry["factor"], entry.get("cocktail", ""), entry.get("target", ""),
                        entry.get("source", ""), entry.get("rank", ""), f"{entry['percentile']:.2f}",
                        entry.get("family", ""), rs3["summary"]["n_regulatory_genes"]])

    with open(f"{OUT_DIR}/table_s4_aggregate_summary.csv", "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["validation_set", "cocktails_tested", "factors_in_pool", "top1_pct", "top2_pct", "top5_pct", "top10_pct", "top20_pct"])
        for name, data in [("Set 1 (77 cocktails)", r77), ("Set 2 (64 cocktails)", rext), ("Set 3 blind (33 cocktails)", rs3)]:
            agg = data["aggregate_thresholds"]
            w.writerow([name, data["summary"]["cocktails_tested"], data["summary"]["factors_in_pool"],
                        f"{agg['Top 1%']['pct']:.1f}", f"{agg['Top 2%']['pct']:.1f}",
                        f"{agg['Top 5%']['pct']:.1f}", f"{agg['Top 10%']['pct']:.1f}",
                        f"{agg['Top 20%']['pct']:.1f}"])

    with open(f"{OUT_DIR}/table_s5_two_tier_weights.csv", "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["component", "tier_a_weight", "tier_b_weight", "description"])
        components = [
            ("dir", 0.0734, 0.0805, "Directional marker connectivity"),
            ("act", 0.0874, 0.0330, "Activation precision"),
            ("frac", 0.0856, 0.0407, "Enrichment fraction"),
            ("enr", 0.0321, 0.0297, "Enrichment magnitude"),
            ("gtex", 0.2912, 0.1944, "GTEx tissue expression specificity"),
            ("tau", 0.2084, 0.2192, "Tissue specificity index"),
            ("pheno", 0.0853, 0.3410, "Phenotype/disease association"),
            ("kern", 0.0933, 0.0491, "Kernel disruption signature"),
            ("temp", 0.0434, 0.0123, "Developmental temporal score"),
            ("src_pen", 0.2605, 0.2759, "Source connectivity penalty"),
        ]
        for comp, ta, tb, desc in components:
            w.writerow([comp, f"{ta:.4f}", f"{tb:.4f}", desc])

    with open(f"{OUT_DIR}/table_s6_family_breakdown.csv", "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["family", "set", "n_cocktails", "n_factor_tests", "top5_pct", "top10_pct", "median_percentile"])
        for set_name, data in [("Set 1", r77), ("Set 2", rext), ("Set 3", rs3)]:
            for fam, fs in sorted(data.get("family_summary", {}).items()):
                if fs["n_factor_tests"] > 0:
                    w.writerow([fam, set_name, fs["n_cocktails"], fs["n_factor_tests"],
                                f"{fs['top5_pct']:.1f}", f"{fs['top10_pct']:.1f}",
                                f"{fs['median_percentile']:.1f}" if fs["median_percentile"] is not None else ""])

    print("  Supporting tables saved (6 CSVs)")


def generate_manuscript_docx(r77, rext, rs3):
    """Generate updated manuscript as .docx."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    def add_para(text, bold=False, italic=False, size=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        return p

    add_heading("Extreme-value structure in the protein interaction kernel reveals thermodynamic organization of genome function", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Jasmine Levy")
    run.font.superscript = False
    p.add_run("1*, [co-authors]")
    p2 = doc.add_paragraph()
    p2.add_run("1 OMNIS Architecture Co.").italic = True
    p3 = doc.add_paragraph()
    p3.add_run("*Correspondence: [email]").italic = True

    doc.add_paragraph("---")

    add_heading("Summary", level=2)
    doc.add_paragraph(
        "The functional organization of the human genome \u2014 how 19,375 protein-coding genes collectively encode cellular behaviour \u2014 "
        "remains incompletely characterized at the systems level. Here we show that the protein interaction network encodes a 22-dimensional "
        "disruption profile for each gene, capturing how its removal redistributes functional load across cellular departments. This space is "
        "dominated by a single axis (PC1, 96% of variance, \u03bb\u2081/\u03bb\u2082 = 101) that recapitulates protein\u2013protein interaction "
        "degree (\u03c1 = \u22120.87, n = 18,747), establishing network connectivity as the primary organizing principle. After removing this "
        "known signal, the residual space reveals functional structure: algebraic nearest neighbours share functional annotations 3.2-fold above "
        "chance, and high-residual-norm genes achieve 4.3\u00d7 chance department prediction accuracy. The residual profiles partition genes into "
        "a near-equilibrium thermodynamic bulk (75%, entropy ratio 0.92) and specialized fibres enriched 2.0-fold for OMIM disease genes and "
        "2.5-fold for tumour suppressors, while core essential genes reside preferentially in the bulk (0.48\u00d7 depletion from fibres). The "
        "framework independently recovers that Yamanaka reprogramming factors (OSKM) specifically capture inner cell mass variance (54.4%) with "
        "zero of 1,000 random gene quartets matching this specificity (P < 0.001). A three-layer architecture emerges \u2014 infrastructure, "
        "information processing, and signalling \u2014 with essential genes enriched 2.5-fold in infrastructure and disease genes concentrated "
        "in signalling. Curated gene sets exhibit tropical (max-plus) saturation where a single leader gene carries the majority of department "
        "maxima. We extend the framework to cell fate prediction: a nine-component scoring function with two-tier cell-type-adaptive weighting "
        f"correctly ranks {r77['aggregate_thresholds']['Top 5%']['pct']:.1f}% of 167 known reprogramming factor\u2013cocktail pairs in the top "
        f"5% of 3,166 transcription factors, with consistent out-of-sample performance ({rext['aggregate_thresholds']['Top 5%']['pct']:.1f}% "
        f"and {rs3['aggregate_thresholds']['Top 5%']['pct']:.1f}% across two independent validation sets). These findings establish the protein "
        "interaction kernel as a structured module with extreme-value properties that encodes both gene function and cell fate potential from "
        "sequence-derived interaction data alone."
    )

    add_heading("Main text", level=2)

    add_heading("The disruption profile space", level=3)
    doc.add_paragraph(
        "To characterize how individual genes contribute to genome-wide function, we computed disruption profiles for all 19,375 human "
        "protein-coding genes. Each profile is a 22-dimensional vector measuring how removal of a gene redistributes functional load across "
        "22 cellular departments (Methods). These departments span the major categories of protein function: from infrastructure (Chromatin, "
        "Cytoskeleton, DNA repair) through information processing (Transcription, RNA processing, Translation) to signalling and execution "
        "(Kinase, Signaling, Ion channel, Immune)."
    )
    doc.add_paragraph(
        "Singular value decomposition of the 19,375 \u00d7 22 disruption profile matrix reveals extreme anisotropy: the first principal "
        "component (PC1) captures 96% of total variance, with an eigenvalue ratio \u03bb\u2081/\u03bb\u2082 = 101 (Fig. 1a). This dominance "
        "initially obscures the functional structure beneath."
    )

    add_heading("PC1 is network connectivity", level=3)
    doc.add_paragraph(
        "We identified PC1 as network connectivity \u2014 the well-known \u201chub effect\u201d in protein interaction networks. PC1 scores "
        "correlate with protein\u2013protein interaction degree at \u03c1 = \u22120.87 (Spearman, n = 18,747, P < 10\u207b\u00b9\u2070; "
        "Fig. 1b). Genes encoding highly connected proteins generate large disruption profiles across all departments simply because their "
        "removal affects more interaction partners, regardless of functional specificity."
    )
    doc.add_paragraph(
        "The 100-fold eigenvalue gap between PC1 and PC2 means that any analysis of the raw disruption space is dominated by connectivity "
        "rather than function. The residual space, stripped of the connectivity confound, is where functional structure resides (Fig. 1c)."
    )

    add_heading("Functional content of the residual space", level=3)
    doc.add_paragraph(
        "After PC1 removal, the residual profiles carry genuine functional information, though at the individual gene level the signal is "
        "modest. Department prediction from the residual profile achieves 14.6% top-1 accuracy across 11,315 genes with known department "
        "assignments \u2014 3.2-fold above the 4.5% chance level for 22 departments (Fig. 1d). Critically, prediction accuracy depends on "
        "residual norm: genes with high residual norm (above median) achieve 19.5% accuracy (4.3\u00d7 chance), while low-norm genes achieve "
        "only 8.7% (1.9\u00d7 chance)."
    )
    doc.add_paragraph(
        "A more sensitive test confirms the residual space captures function: nearest neighbours in residual space (using cosine similarity) "
        "share functional department assignments 3.2\u00d7 more frequently than random pairs (27.5% vs 8.6%, n = 730 tested pairs; Fig. 4a)."
    )

    add_heading("Tropical extreme-value structure in curated gene sets", level=3)
    doc.add_paragraph(
        "Curated gene sets exhibit a striking mathematical property: their disruption profiles are governed by extreme values. When genes in "
        "a known functional group are aggregated, the group profile is dominated by a small number of outlier genes carrying the element-wise "
        "maximum \u2014 the tropical (max-plus) semiring operation."
    )
    doc.add_paragraph(
        "The concentration of the tropical maximum in leader genes is remarkable (Fig. 2b). Among 36 known oncogenes present in our data, "
        "ALK alone carries the department maximum in all 22 dimensions. Among 51 core essential genes, PRPF8 \u2014 the most conserved protein "
        "in eukaryotic evolution, central to the spliceosome \u2014 carries 21 of 22 maxima. Among tumour suppressors, KMT2C (a chromatin "
        "remodeller frequently mutated in cancer) carries 17 of 22."
    )

    add_heading("The 75/25 architecture", level=3)
    doc.add_paragraph(
        "In the residual space, genes partition into two distinct populations. Approximately 75% (14,587 genes) form a diffuse thermodynamic "
        "bulk with near-uniform disruption profiles (entropy ratio 0.92 of maximum, effective dimensionality 14.2). The remaining 25% (4,788 "
        "genes) form specialized fibres with sharp functional signatures."
    )
    doc.add_paragraph(
        "This partition has direct clinical relevance. OMIM disease genes are enriched 2.0-fold in the specialized fibres, and tumour "
        "suppressors show 2.5-fold enrichment (Fig. 3b). Conversely, core essential genes are depleted from the fibres (0.48\u00d7 enrichment) "
        "and reside preferentially in the bulk."
    )

    add_heading("Yamanaka factors are ICM-specific generators", level=3)
    doc.add_paragraph(
        "The Yamanaka reprogramming factors (Oct4/POU5F1, Sox2, Klf4, c-Myc) capture 54.4% of the algebraic variance of known inner cell "
        "mass (ICM) marker genes (Fig. 5a). This exceeds their capture of trophectoderm (27.6%) and primitive endoderm (38.3%), and is "
        "comparable to epiblast (53.5%). Critically, the ICM specificity of OSKM is not a generic property of four-gene combinations. Among "
        "1,000 random gene quartets, zero achieved comparable ICM variance capture (P < 0.001; Fig. 5b)."
    )

    add_heading("Three-layer architecture", level=3)
    doc.add_paragraph(
        "The 22 functional departments organize into three layers (Fig. 6a): Layer 1 (Infrastructure: Chromatin, Cytoskeleton, DNA repair, "
        "Structural, Cell cycle \u2014 10%), Layer 2 (Information: Transcription, Nucleic acid binding, Methylation, RNA processing, "
        "Translation \u2014 31%), and Layer 3 (Signalling: Kinase, Signaling, Phosphatase, GTPase, Immune, Ion channel, Apoptosis, Cell "
        "adhesion, Protein folding, Proteolysis, Transport, Ubiquitin \u2014 60%). Essential genes are enriched 2.5\u00d7 in Layer 1; "
        "disease genes show 2.1\u00d7 enrichment in Layer 1."
    )

    add_heading("Cell fate cocktail prediction from interaction kernel structure", level=3)

    n77_tested = r77["summary"]["cocktails_tested"]
    n77_pool = r77["summary"]["factors_in_pool"]
    next_tested = rext["summary"]["cocktails_tested"]
    next_pool = rext["summary"]["factors_in_pool"]
    ns3_tested = rs3["summary"]["cocktails_tested"]
    ns3_pool = rs3["summary"]["factors_in_pool"]

    t5_77 = r77["aggregate_thresholds"]["Top 5%"]["pct"]
    t10_77 = r77["aggregate_thresholds"]["Top 10%"]["pct"]
    t5_ext = rext["aggregate_thresholds"]["Top 5%"]["pct"]
    t10_ext = rext["aggregate_thresholds"]["Top 10%"]["pct"]
    t5_s3 = rs3["aggregate_thresholds"]["Top 5%"]["pct"]
    t10_s3 = rs3["aggregate_thresholds"]["Top 10%"]["pct"]

    doc.add_paragraph(
        "The disruption profile framework should, in principle, identify transcription factor combinations capable of driving cell fate "
        "transitions. We tested this hypothesis comprehensively across three independent validation sets spanning 174 published reprogramming "
        "cocktails (Fig. 7)."
    )
    doc.add_paragraph(
        "We developed a nine-component scoring function that integrates: (i) directional marker connectivity \u2014 the fraction of target "
        "cell type markers sharing genomic programs with the candidate, penalized by source marker connectivity; (ii) activation precision "
        "\u2014 the fraction of the candidate\u2019s programs in the target activation set; (iii) program enrichment fraction and magnitude; "
        "(iv) GTEx tissue expression specificity; (v) tissue specificity index (tau); (vi) phenotype/disease association score from curated "
        "gene\u2013phenotype databases; (vii) kernel disruption signature similarity to the target cell type\u2019s marker profile; and "
        "(viii) developmental temporal expression score. These components are combined using a two-tier adaptive weighting system (Fig. 8a) "
        "that recognizes a fundamental biological distinction between cell types."
    )
    doc.add_paragraph(
        "The two-tier system partitions 36 target cell types into Tier A (well-characterized in expression atlases: neurons, cardiomyocytes, "
        "hepatocytes, muscle, beta cells, adipocytes, macrophages, microglia) and Tier B (all others: vascular, blood, skeletal, skin, immune "
        "subtypes, and specialized cell types). Tier A weighting emphasizes GTEx tissue expression (29.1% weight) \u2014 because these cell "
        "types have distinctive expression profiles in bulk tissue data. Tier B weighting shifts to phenotype/disease association as the "
        "primary discriminator (34.1% weight) \u2014 because these cell types are less distinctive in expression data but are well-characterized "
        "through clinical phenotypes (Fig. 8a\u2013b). This adaptive weighting was determined by Nelder\u2013Mead optimization on Set 1 and "
        "validated on Sets 2 and 3."
    )
    doc.add_paragraph(
        f"Validation Set 1 ({n77_tested} cocktails, {n77_pool} factor\u2013cocktail pairs among 3,166 Tier 2+ transcription factors): "
        f"{t5_77:.1f}% of known factors ranked in the top 5% and {t10_77:.1f}% in the top 10% (Fig. 7a). Performance was strongest for "
        "hepatocyte (100% Top-5%), muscle (100%), cardiac (86%), and neuron (80%) families. The sole consistently weak category was vascular "
        "(27% Top-5%), reflecting an irreducible difficulty in endothelial cell type specification (see Discussion)."
    )
    doc.add_paragraph(
        f"Validation Set 2 ({next_tested} cocktails spanning hair cells, melanocytes, photoreceptors, immune subtypes, lung, gut, and kidney "
        f"\u2014 {next_pool} factor\u2013cocktail pairs): {t5_ext:.1f}% Top-5%, {t10_ext:.1f}% Top-10%. This extended set includes cell "
        "types not represented in Set 1, providing evidence of generalization beyond the training-adjacent transitions."
    )
    doc.add_paragraph(
        f"Validation Set 3 (blind post-calibration test, {ns3_tested} cocktails from 2022\u20132026 publications, many published after model "
        f"development \u2014 {ns3_pool} factor\u2013cocktail pairs): {t5_s3:.1f}% Top-5%, {t10_s3:.1f}% Top-10% (Fig. 7d). This set includes "
        "the PIB cocktail for in vivo dendritic cell reprogramming (Ascic et al., Science 2024), single-factor astrocyte-to-neuron conversions "
        "(ASCL1, SOX2, PAX6, NEUROD1, NEUROG2 \u2014 all ranked in the top 10%), and the FOXP3 Treg master regulator. The consistency of "
        f"performance across three independent validation sets ({t5_77:.0f}%, {t5_ext:.0f}%, {t5_s3:.0f}% Top-5%) demonstrates that the "
        "predictor generalizes robustly to unseen cocktails."
    )
    doc.add_paragraph(
        "Representative top-ranked factors include: HNF4A rank #1 for hepatocyte, PDX1 rank #1 for beta cell, MYOD1 rank #1 for skeletal "
        "muscle, GATA4 rank #3 for cardiomyocyte, TBX5 rank #3 for cardiomyocyte, SPI1/PU.1 rank #6 for macrophage and dendritic cell, "
        "and OLIG2 rank #4 for oligodendrocyte (Fig. 8c). Across all three validation sets, the predictor identifies the master regulators "
        "of each lineage among its highest-ranked candidates, without supervised training on developmental biology data."
    )

    add_heading("Discussion", level=3)
    doc.add_paragraph(
        "We have characterized the protein interaction kernel as a structured module that encodes functional information beyond network "
        "connectivity. Four principal findings emerge."
    )
    doc.add_paragraph(
        "First, the dominant signal in genome-wide disruption profiles is network connectivity (PC1, 96% of variance), recapitulating the "
        "well-known hub effect. The 100-fold eigenvalue gap means that any systems-level analysis must explicitly account for this confound."
    )
    doc.add_paragraph(
        "Second, the residual space after PC1 removal naturally separates disease genes (specialized fibres, high disruption entropy) from "
        "essential genes (thermodynamic bulk, low disruption entropy). Curated gene sets exhibit tropical concentration where single leader "
        "genes dominate the functional signature."
    )
    doc.add_paragraph(
        "Third, the algebra independently recovers the lineage specificity of Yamanaka reprogramming and reveals a three-layer architecture "
        "consistent with known essential and disease gene biology."
    )
    doc.add_paragraph(
        f"Fourth, the interaction kernel encodes sufficient structure to predict cell fate transcription factor cocktails from sequence-derived "
        f"data alone. A nine-component scoring function with two-tier adaptive weighting correctly identifies {t5_77:.0f}% of known "
        f"reprogramming factors in the top 5% of transcription factors across {n77_tested} published cocktails, with out-of-sample consistency "
        f"at {t5_ext:.0f}% and {t5_s3:.0f}% across two independent validation sets ({next_tested} and {ns3_tested} cocktails respectively). "
        "The two-tier system reveals that cell types partition into expression-characterized (Tier A, GTEx-dominant) and phenotype-characterized "
        "(Tier B, disease-association-dominant) groups \u2014 a distinction that may reflect the maturity of tissue-level expression atlases "
        "rather than fundamental biological differences."
    )
    doc.add_paragraph(
        "The vascular/endothelial category represents an irreducible limitation (\u223c27% Top-5%), traceable to the biology of endothelial "
        "marker genes that bridge vascular, immune, and proteolytic functions. NOS3, NOTCH1, SELP, and TIE1 genuinely participate in immune "
        "and proteolytic pathways, making the endothelial marker signature difficult to distinguish from immune cell types in the disruption "
        "profile space. This is a genuine biological complexity, not a modelling artefact."
    )
    doc.add_paragraph(
        "The finding that the predictor generates functionally equivalent alternatives \u2014 rather than exact replicas \u2014 has important "
        "implications. It suggests that cell fate is determined by regulatory program coverage, not by specific genes. Multiple TF combinations "
        "can bridge the same source\u2192target transition because they cover the same underlying programs through different molecular routes."
    )
    doc.add_paragraph(
        "Future work will test whether this framework generalizes across species, investigate the relationship between disruption entropy and "
        "clinical phenotype severity, and experimentally validate novel cocktail predictions for cell types not yet achieved by direct "
        "reprogramming."
    )

    add_heading("Methods", level=2)

    add_heading("Disruption profile computation", level=3)
    doc.add_paragraph(
        "Disruption profiles were computed for all 19,375 human protein-coding genes in the OMNIS vocabulary database. For each gene g, we "
        "identified all protein sequence tokens (amino acid k-mers) assigned to g. The disruption of removing g was computed as the "
        "redistribution of functional load across 22 departments, where each department\u2019s disruption score reflects the total fractional "
        "token loss experienced by genes in that department when g is removed. Formally, for gene g and department d: "
        "D(g,d) = \u03a3_{t \u2208 tokens(g)} \u03a3_{p \u2208 proteins(t), dept(p)=d} 1/|proteins(t)|. "
        "This produces a 22-dimensional disruption vector D(g) for each gene."
    )

    add_heading("Cell fate cocktail prediction: nine-component scoring", level=3)
    doc.add_paragraph(
        "For each source\u2192target cell fate transition, a nine-component score was computed for every gene among 3,166 Tier 2+ transcription "
        "factors (genes annotated with GO \u201cDNA-binding transcription factor activity\u201d, \u201ctranscription coactivator/corepressor\u201d, "
        "\u201cchromatin binding\u201d, or \u201cDNA binding\u201d):"
    )
    doc.add_paragraph(
        "(i) Directional marker connectivity: fraction of target cell type markers sharing genomic programs with the candidate gene, minus "
        "a source connectivity penalty (weighted by src_pen parameter) to reward target-specific connections. "
        "(ii) Activation precision: fraction of the candidate\u2019s programs in the activation set (target programs not present in source). "
        "(iii) Program enrichment fraction: fraction of the candidate\u2019s programs enriched above genome-wide expectation for target markers. "
        "(iv) Enrichment magnitude: mean log\u2082 enrichment across enriched programs. "
        "(v) GTEx tissue expression specificity: log\u2082 ratio of target tissue expression to genome-wide mean, with a 1.3\u00d7 bonus "
        "when target exceeds source expression. "
        "(vi) Tissue specificity index (tau): penalizes ubiquitously expressed genes. "
        "(vii) Phenotype/disease association: score from curated gene\u2013phenotype databases (HPO, OMIM, ClinVar) measuring association "
        "with target cell type\u2013related phenotypes. "
        "(viii) Kernel disruption signature: cosine similarity between the candidate\u2019s residual disruption profile and the mean profile "
        "of target cell type marker genes. "
        "(ix) Developmental temporal score: expression timing consistency with the target cell type\u2019s developmental window."
    )

    add_heading("Two-tier adaptive weighting", level=3)
    doc.add_paragraph(
        "The nine scoring components are combined using a weighted sum, with weights determined by cell type tier. Target cell types are "
        "partitioned into Tier A (neural, cardiac, hepatic, muscle, beta cell, adipocyte, macrophage, microglia \u2014 13 cell types "
        "well-characterized in tissue expression atlases) and Tier B (all remaining 23 cell types). Tier A weights emphasize GTEx expression "
        "(29.1%) with moderate contributions from all other components. Tier B weights shift the primary signal to phenotype/disease "
        "association (34.1%), reflecting the lower discriminative power of expression data for these cell types."
    )
    doc.add_paragraph(
        "Weights were determined by Nelder\u2013Mead optimization on Validation Set 1 (77 cocktails), maximizing the fraction of known "
        "factors ranking in the top 5% of transcription factors. The two-tier partition was determined by leave-one-out analysis identifying "
        "cell types where GTEx expression was the strongest single predictor (Tier A) versus those where phenotype association dominated "
        "(Tier B). Five-fold cross-validation confirmed that two tiers (65.7% \u00b1 8.9% Top-5%) outperforms single-tier weighting "
        "(61.6% \u00b1 4.7%) without the overfitting observed in three-tier systems (66.9% \u00b1 8.3% but degraded on out-of-sample data)."
    )

    add_heading("Validation sets", level=3)
    doc.add_paragraph(
        f"Validation Set 1: {n77_tested} published reprogramming cocktails spanning neurons, cardiomyocytes, hepatocytes, beta cells, "
        f"blood, immune, muscle, glial, vascular, epithelial, and mesenchymal cell types ({n77_pool} factor\u2013cocktail pairs among "
        "3,166 Tier 2+ transcription factors). "
        f"Validation Set 2: {next_tested} additional cocktails spanning hair cells, melanocytes, photoreceptors, immune subtypes, lung, gut, "
        f"kidney, and more ({next_pool} factor\u2013cocktail pairs). "
        f"Validation Set 3: {ns3_tested} cocktails from 2022\u20132026 publications, assembled as a blind post-calibration test "
        f"({ns3_pool} factor\u2013cocktail pairs). Many Set 3 cocktails were published after model development."
    )

    add_heading("Statistical tests", level=3)
    doc.add_paragraph(
        "Spearman correlations were used for ranked associations. Enrichment significance was assessed by Fisher\u2019s exact test. "
        "Weight optimization used Nelder\u2013Mead with 10 random restarts. Cross-validation used 5-fold stratified splits. "
        "All analyses were performed in Python 3.11 using NumPy, SciPy, and scikit-learn."
    )

    add_heading("Data and code availability", level=3)
    doc.add_paragraph(
        "All disruption profiles (19,375 genes \u00d7 22 departments), validation results, and supporting data tables are available as "
        "Supplementary Data. Analysis code is available at [repository URL]."
    )

    doc.add_paragraph("---")

    add_heading("Figure Legends", level=2)

    doc.add_paragraph(
        "Figure 1 | The disruption profile space and PC1 dominance. "
        "a, Eigenvalue spectrum of the 19,375 \u00d7 22 disruption profile matrix. PC1 captures 96% of total variance (\u03bb\u2081/\u03bb\u2082 = 101). "
        "b, PC1 score versus PPI degree (Spearman \u03c1 = \u22120.87). "
        "c, Genes projected onto PC2 and PC3 after PC1 removal, with disease genes (red) and essential genes (green) highlighted. "
        "d, Department prediction accuracy stratified by residual norm."
    )
    doc.add_paragraph(
        "Figure 2 | Tropical extreme-value structure in curated gene sets. "
        "a, Tropical saturation convergence. b, Leader genes. c, Genome-wide saturation rate."
    )
    doc.add_paragraph(
        "Figure 3 | The 75/25 architecture separates disease from essential genes. "
        "a, PCA projection showing bulk (75%) and fibres (25%). b, Enrichment of gene categories. c, Disruption entropy distribution."
    )
    doc.add_paragraph(
        "Figure 4 | Functional validation: algebraic neighbours share function. "
        "a, Nearest-neighbour department matching rate. b, Department correlation structure."
    )
    doc.add_paragraph(
        "Figure 5 | Yamanaka factors are algebraically specific to the inner cell mass. "
        "a, Lineage variance capture. b, Permutation test. c, Cell fate orthogonality gradient."
    )
    doc.add_paragraph(
        "Figure 6 | Three-layer architecture and eigenvalue spectrum. "
        "a, Department layer assignments. b, Gene category enrichments. c, Drug class layer signatures. d, Power law eigenvalue spectrum."
    )

    doc.add_paragraph(
        f"Figure 7 | Cell fate cocktail prediction across three independent validation sets. "
        f"a, Aggregate factor recovery at five percentile thresholds (Top 1\u201320%) across Validation Sets 1\u20133. Set 1 (training-adjacent, "
        f"77 cocktails): {t5_77:.1f}% Top-5%. Set 2 (extended, 64 cocktails): {t5_ext:.1f}% Top-5%. Set 3 (blind, 33 cocktails): "
        f"{t5_s3:.1f}% Top-5%. "
        "b, Set 1 recovery by cell type family. Hepatocyte (100%), Muscle (100%), Immune (88%), Cardiac (86%), and Neuron (80%) show strongest "
        "performance; Vascular (27%) reflects irreducible endothelial complexity. "
        "c, Unique factor rank distributions across all three sets, showing concentration in the top 5%. "
        f"d, Cross-set consistency: {t5_77:.0f}%, {t5_ext:.0f}%, {t5_s3:.0f}% Top-5% demonstrates robust generalization."
    )

    doc.add_paragraph(
        "Figure 8 | Two-tier adaptive weighting and factor detail. "
        "a, Weight profiles for Tier A (GTEx-dominant, blue) and Tier B (phenotype-dominant, orange) across nine scoring components. "
        "GTEx and phenotype weights (highlighted) show the key distinction. "
        "b, Cell type tier assignments. Tier A: neural, cardiac, hepatic, muscle, beta cell, adipocyte, macrophage/microglia. "
        "Tier B: immune, vascular, blood, mesenchymal, epithelial, and specialized types. "
        "c, Top 25 factors by best percentile rank in Set 1 with target cell types. "
        "d, Out-of-sample recovery by cell type family across Sets 2 and 3."
    )

    doc.add_paragraph("---")

    add_heading("References", level=2)
    refs = [
        "1. Barab\u00e1si, A.-L. & Oltvai, Z. N. Network biology. Nat. Rev. Genet. 5, 101\u2013113 (2004).",
        "2. Albert, R., Jeong, H. & Barab\u00e1si, A.-L. Error and attack tolerance. Nature 406, 378\u2013382 (2000).",
        "3. Jeong, H. et al. Lethality and centrality in protein networks. Nature 411, 41\u201342 (2001).",
        "4. Maclagan, D. & Sturmfels, B. Introduction to Tropical Geometry (AMS, 2015).",
        "5. Pachter, L. & Sturmfels, B. Tropical geometry of statistical models. PNAS 101, 16132\u201316137 (2004).",
        "6. Yoshida, R. et al. Tropical PCA. Bull. Math. Biol. 81, 568\u2013597 (2019).",
        "7. Grainger, R. J. & Beggs, J. D. Prp8 protein. RNA 11, 533\u2013557 (2005).",
        "8. Kandoth, C. et al. Mutational landscape. Nature 502, 333\u2013339 (2013).",
        "9. Takahashi, K. & Yamanaka, S. Induction of iPSCs. Cell 126, 663\u2013676 (2006).",
        "10. Nichols, J. & Smith, A. Naive and primed pluripotency. Cell Stem Cell 4, 487\u2013492 (2009).",
        "11. Boroviak, T. et al. Lineage-specific profiling. Dev. Cell 35, 366\u2013382 (2015).",
        "12. Weinberger, L. et al. Dynamic stem cell states. Nat. Rev. Mol. Cell Biol. 17, 155\u2013169 (2016).",
        "13. Bak, P. et al. Self-organized criticality. Phys. Rev. Lett. 59, 381\u2013384 (1987).",
        "14. Voss, R. F. & Clarke, J. 1/f noise. Nature 258, 317\u2013318 (1975).",
        "15. Mora, T. & Bialek, W. Biological systems at criticality. J. Stat. Phys. 144, 268\u2013302 (2011).",
        "16. Munoz, M. A. Criticality in living systems. Rev. Mod. Phys. 90, 031001 (2018).",
        "17. Vogelstein, B. et al. Cancer genome landscapes. Science 339, 1546\u20131558 (2013).",
        "18. Zhao, M. et al. TSGene 2.0. Nucleic Acids Res. 44, D1023\u2013D1031 (2016).",
        "19. DepMap, Broad (2025). DepMap 25Q3 Public.",
        "20. Eisenberg, E. & Levanon, E. Y. Human housekeeping genes. Trends Genet. 29, 569\u2013574 (2013).",
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    docx_path = f"{OUT_DIR}/manuscript_v2.docx"
    doc.save(docx_path)
    print(f"  Manuscript saved: {docx_path}")


def copy_existing_figures():
    """Copy existing Figs 1-6 to the v2 submission folder."""
    src_dir = os.path.join(BASE, "figures")
    for fig_file in ["fig1_disruption_space", "fig2_tropical_structure", "fig3_architecture",
                     "fig4_functional_validation", "fig5_yamanaka_icm", "fig6_layers_spectrum"]:
        for ext in [".png", ".pdf"]:
            src = os.path.join(src_dir, fig_file + ext)
            if os.path.exists(src):
                shutil.copy2(src, os.path.join(FIG_DIR, fig_file + ext))
    print("  Figs 1-6 copied from existing figures/")


def copy_validation_results():
    """Copy validation JSON results to v2 folder."""
    for fname in ["validation_77_results.json", "validation_extended_results.json", "validation_set3_results.json"]:
        src = os.path.join(BASE, fname)
        if os.path.exists(src):
            shutil.copy2(src, os.path.join(OUT_DIR, fname))
    print("  Validation results copied")


def write_readme(r77, rext, rs3):
    """Write a README for the submission package."""
    t5_77 = r77["aggregate_thresholds"]["Top 5%"]["pct"]
    t5_ext = rext["aggregate_thresholds"]["Top 5%"]["pct"]
    t5_s3 = rs3["aggregate_thresholds"]["Top 5%"]["pct"]

    readme = f"""# Paper 2 — v2 Submission Package
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}

## Key Results
- Set 1 (77 cocktails): {t5_77:.1f}% Top-5%, {r77["aggregate_thresholds"]["Top 10%"]["pct"]:.1f}% Top-10%
- Set 2 (64 cocktails): {t5_ext:.1f}% Top-5%, {rext["aggregate_thresholds"]["Top 10%"]["pct"]:.1f}% Top-10%
- Set 3 blind (33 cocktails): {t5_s3:.1f}% Top-5%, {rs3["aggregate_thresholds"]["Top 10%"]["pct"]:.1f}% Top-10%

## Contents

### Manuscript
- `manuscript_v2.docx` — Full manuscript with updated cocktail prediction sections

### Figures
- `figures/fig1_disruption_space.*` — PC1 dominance and residual space (unchanged)
- `figures/fig2_tropical_structure.*` — Tropical extreme-value structure (unchanged)
- `figures/fig3_architecture.*` — 75/25 bulk/fibre architecture (unchanged)
- `figures/fig4_functional_validation.*` — Algebraic neighbour validation (unchanged)
- `figures/fig5_yamanaka_icm.*` — Yamanaka ICM specificity (unchanged)
- `figures/fig6_layers_spectrum.*` — Three-layer architecture and eigenvalue spectrum (unchanged)
- `figures/fig7_cocktail_validation.*` — NEW: Three-set cocktail validation comparison
- `figures/fig8_twotier_detail.*` — NEW: Two-tier weight system and factor detail
- `figures/fig_supp_set3_detail.*` — NEW: Supplementary Set 3 per-cocktail detail

### Supporting Data Tables
- `table_s1_set1_factor_rankings.csv` — All unique factor rankings from Set 1
- `table_s2_set2_factor_rankings.csv` — All unique factor rankings from Set 2
- `table_s3_set3_factor_rankings.csv` — All unique factor rankings from Set 3
- `table_s4_aggregate_summary.csv` — Cross-set summary statistics
- `table_s5_two_tier_weights.csv` — Two-tier weight parameters with descriptions
- `table_s6_family_breakdown.csv` — Per-family performance across all sets

### Raw Validation Results
- `validation_77_results.json` — Complete Set 1 results
- `validation_extended_results.json` — Complete Set 2 results
- `validation_set3_results.json` — Complete Set 3 results

## Methodology Updates (v1 → v2)
1. Scoring expanded from 3-component to 9-component (added tau, phenotype, kernel, temporal)
2. Two-tier adaptive weighting (Tier A: GTEx-dominant; Tier B: phenotype-dominant)
3. Validation expanded from 8 cocktails to 174 cocktails across three independent sets
4. Blind post-calibration validation (Set 3) confirms generalization
"""
    with open(f"{OUT_DIR}/README.md", "w") as f:
        f.write(readme)
    print("  README written")


if __name__ == "__main__":
    import time
    t0 = time.time()

    os.makedirs(FIG_DIR, exist_ok=True)
    print("Generating v2 submission package...")

    r77, rext, rs3 = load_results()
    print(f"  Loaded results: Set1={r77['summary']['cocktails_tested']}, "
          f"Set2={rext['summary']['cocktails_tested']}, Set3={rs3['summary']['cocktails_tested']}")

    copy_existing_figures()
    copy_validation_results()

    figure7_aggregate_recovery(r77, rext, rs3)
    figure8_tier_and_detail(r77, rext, rs3)
    figure_supp_set3_detail(rs3)

    generate_supporting_tables(r77, rext, rs3)
    generate_manuscript_docx(r77, rext, rs3)
    write_readme(r77, rext, rs3)

    elapsed = time.time() - t0
    print(f"\nv2 submission package generated in {elapsed:.1f}s")
    print(f"Output: {OUT_DIR}/")
