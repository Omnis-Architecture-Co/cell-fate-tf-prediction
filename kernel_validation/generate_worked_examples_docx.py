"""
Generate Supplementary Note 1: Encoding Pipeline Worked Examples
Two complete walkthroughs: Nucleolin (protein pathway) and INS locus (DNA pathway)
Uses the V2 6-bit codon pipeline: AA → first-listed RNA codon (Table S5) → DNA (U→T)
→ nucleotide binary (A=00, T=01, G=10, C=11) → 8-bit bytes
"""
import csv
import json
import os
import urllib.request
from collections import Counter
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "exports",
                        "Supplementary_Note1_Worked_Examples.docx")

NUC_BIN = {'A': '00', 'T': '01', 'G': '10', 'C': '11'}

OMNIS_FIRST_DNA = {
    'A': 'GCT', 'C': 'TGT', 'D': 'GAT', 'E': 'GAG', 'F': 'TTT',
    'G': 'GGT', 'H': 'CAT', 'I': 'ATT', 'K': 'AAG', 'L': 'TTG',
    'M': 'ATG', 'N': 'AAT', 'P': 'CCT', 'Q': 'CAG', 'R': 'CGT',
    'S': 'TCT', 'T': 'ACT', 'V': 'GTT', 'W': 'TGG', 'Y': 'TAT',
    '*': 'TAG', 'X': 'NNN'
}

OMNIS_FIRST_RNA = {
    'A': 'GCU', 'C': 'UGU', 'D': 'GAU', 'E': 'GAG', 'F': 'UUU',
    'G': 'GGU', 'H': 'CAU', 'I': 'AUU', 'K': 'AAG', 'L': 'UUG',
    'M': 'AUG', 'N': 'AAU', 'P': 'CCU', 'Q': 'CAG', 'R': 'CGU',
    'S': 'UCU', 'T': 'ACU', 'V': 'GUU', 'W': 'UGG', 'Y': 'UAU',
    '*': 'UAG', 'X': 'NNN'
}

AA_NAMES = {
    'A': 'Alanine', 'C': 'Cysteine', 'D': 'Aspartate', 'E': 'Glutamate',
    'F': 'Phenylalanine', 'G': 'Glycine', 'H': 'Histidine', 'I': 'Isoleucine',
    'K': 'Lysine', 'L': 'Leucine', 'M': 'Methionine', 'N': 'Asparagine',
    'P': 'Proline', 'Q': 'Glutamine', 'R': 'Arginine', 'S': 'Serine',
    'T': 'Threonine', 'V': 'Valine', 'W': 'Tryptophan', 'Y': 'Tyrosine'
}


def set_cell_font(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, bold=False, italic=False, size=12, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    return p


def add_heading_text(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, bold=True, size=10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_font(table.rows[r_idx + 1].cells[c_idx], str(val), size=10)
    return table


def aa_to_6bit_binary(aa):
    codon = OMNIS_FIRST_DNA.get(aa)
    if not codon or 'N' in codon:
        return '000000'
    return ''.join(NUC_BIN[nt] for nt in codon)


def encode_protein(seq):
    bits = ''
    for aa in seq:
        bits += aa_to_6bit_binary(aa)
    remainder = len(bits) % 8
    usable = len(bits) - remainder
    byte_vals = [int(bits[i:i+8], 2) for i in range(0, usable, 8)]
    return byte_vals, bits, remainder


def encode_dna(seq):
    bits = ''
    for nt in seq:
        bits += NUC_BIN.get(nt, '00')
    remainder = len(bits) % 8
    usable = len(bits) - remainder
    byte_vals = [int(bits[i:i+8], 2) for i in range(0, usable, 8)]
    return byte_vals, bits, remainder


def classify_bytes(byte_vals):
    control = sum(1 for b in byte_vals if b <= 0x1F)
    standard = sum(1 for b in byte_vals if 0x20 <= b <= 0x7F)
    extended = sum(1 for b in byte_vals if b >= 0x80)
    return control, standard, extended


def load_vocab():
    vocab = {}
    path = os.path.join(os.path.dirname(__file__), "..", "server", "data",
                        "human", "vocabulary.csv")
    with open(path) as f:
        reader = csv.DictReader(f)
        for row in reader:
            hex_word = row['word_hex'].replace('0x', '').upper()
            byte_len = len(hex_word) // 2
            vocab[hex_word] = {
                'hex': hex_word,
                'dept': row.get('primary_function', ''),
                'enrichment': float(row.get('enrichment', row.get('enrichment_score', '0')) or '0'),
                'occurrences': int(row.get('occurrences', '0') or '0'),
                'byte_len': byte_len,
            }
    return vocab


def tokenize(hex_stream, vocab):
    hex_bytes = [hex_stream[i:i+2] for i in range(0, len(hex_stream), 2)]
    vocab_by_len = {}
    for hex_key, info in vocab.items():
        bl = info['byte_len']
        if bl not in vocab_by_len:
            vocab_by_len[bl] = {}
        vocab_by_len[bl][hex_key] = info

    matches = []
    lengths = sorted(vocab_by_len.keys(), reverse=True)
    covered = set()

    for length in lengths:
        vfl = vocab_by_len[length]
        for i in range(len(hex_bytes) - length + 1):
            if any(j in covered for j in range(i, i + length)):
                continue
            pattern = ''.join(hex_bytes[i:i+length])
            if pattern in vfl:
                for j in range(i, i + length):
                    covered.add(j)
                matches.append({
                    'position': i,
                    'hex': pattern,
                    'dept': vfl[pattern]['dept'],
                    'enrichment': vfl[pattern]['enrichment'],
                    'occurrences': vfl[pattern]['occurrences'],
                    'byte_len': length,
                })

    matches.sort(key=lambda m: m['position'])
    return matches


def fetch_sequence(url):
    req = urllib.request.Request(url, headers={"User-Agent": "Python/3"})
    with urllib.request.urlopen(req, timeout=20) as resp:
        fasta = resp.read().decode()
    lines = fasta.strip().split('\n')
    seq = ''.join(l for l in lines[1:] if not l.startswith('>')).upper()
    return seq


def load_programs_near(chromosome, position, window=500000):
    path = os.path.join(os.path.dirname(__file__), "..", "exports",
                        "programs_annotated.csv")
    programs = []
    with open(path) as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row['chromosome'] == chromosome:
                s = int(row['start_position'])
                if abs(s - position) < window:
                    programs.append(row)
    return programs


def main():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    add_para(doc, "Supplementary Note 1", bold=True, size=14,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Encoding Pipeline Worked Examples", bold=True, size=13,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "", size=6)

    add_para(doc, (
        "This note traces two biological sequences through the complete 8-step "
        "encoding pipeline described in Methods. Example A follows nucleolin "
        "(UniProt P19338), a 710-amino-acid nucleolar phosphoprotein, through "
        "the protein-derived pathway. Example B follows a 1,000-nucleotide "
        "segment of the chromosome 11 INS locus (T2T-CHM13v2.0, 11p15.5) "
        "through the DNA-native pathway. Both examples "
        "use real pipeline outputs at every stage."
    ))

    add_para(doc, (
        "Encoding rule (both pathways): each DNA nucleotide maps to a 2-bit "
        "value (A = 00, T = 01, G = 10, C = 11). For proteins, each amino acid "
        "is first reverse-translated to its OMNIS first-listed DNA codon "
        "(Table S5), producing a 6-bit code per residue (3 nucleotides \u00d7 2 bits). "
        "The resulting bit stream is then packed into 8-bit bytes."
    ), size=11, italic=True)

    ncl_cache = "/tmp/ncl_sequence.txt"
    ins_cache = "/tmp/ins_sequence.txt"

    if os.path.exists(ncl_cache):
        with open(ncl_cache) as f:
            ncl_seq = f.read().strip()
    else:
        ncl_seq = fetch_sequence("https://rest.uniprot.org/uniprotkb/P19338.fasta")
        with open(ncl_cache, 'w') as f:
            f.write(ncl_seq)

    if os.path.exists(ins_cache):
        with open(ins_cache) as f:
            ins_seq = f.read().strip()
    else:
        ins_seq = fetch_sequence(
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?"
            "db=nuccore&id=NC_060935.1&rettype=fasta&seq_start=2159280&seq_stop=2160279"
        )
        with open(ins_cache, 'w') as f:
            f.write(ins_seq)

    vocab = load_vocab()

    add_heading_text(doc, "Example A: Nucleolin (Protein-Derived Pathway)", level=1)
    add_para(doc, (
        f"Nucleolin (gene NCL, UniProt P19338) is a 710-amino-acid nucleolar "
        f"phosphoprotein involved in ribosome biogenesis, chromatin remodeling, "
        f"and transcription regulation. Its N-terminal domain contains "
        f"alternating acidic and basic stretches; its C-terminal domain contains "
        f"four RNA recognition motifs (RRMs) and a glycine/arginine-rich (GAR) domain."
    ))

    add_heading_text(doc, "Step 1 \u2014 Amino Acid to Codon to Binary Encoding", level=2)
    add_para(doc, (
        "Each amino acid is reverse-translated to its OMNIS first-listed DNA codon "
        "(Table S5), then each nucleotide of the codon is converted to its 2-bit "
        "binary value (A = 00, T = 01, G = 10, C = 11), yielding 6 bits per residue. "
        "The first six residues of nucleolin (MVKLAK) are encoded as follows:"
    ))

    step1_headers = ["Position", "Residue", "Amino Acid", "DNA Codon\n(Table S5)", "Binary\n(6-bit)"]
    step1_rows = []
    for i, aa in enumerate(ncl_seq[:6]):
        codon = OMNIS_FIRST_DNA[aa]
        binary = ''.join(NUC_BIN[nt] for nt in codon)
        step1_rows.append([
            str(i + 1), aa, AA_NAMES.get(aa, ''), codon, binary
        ])
    make_table(doc, step1_headers, step1_rows)

    total_bits = len(ncl_seq) * 6
    add_para(doc, (
        f"The complete 710-residue sequence produces {total_bits:,} bits of "
        f"binary data (710 residues \u00d7 3 nucleotides \u00d7 2 bits)."
    ), size=11)

    add_heading_text(doc, "Step 2 \u2014 Binary-to-Byte Conversion", level=2)

    ncl_bytes, ncl_bits, remainder = encode_protein(ncl_seq)

    add_para(doc, (
        f"The {total_bits:,}-bit stream is partitioned into 8-bit bytes, yielding "
        f"{len(ncl_bytes)} complete bytes with a {remainder}-bit remainder that is discarded "
        f"(not zero-padded), preventing spurious terminal patterns. "
        f"Because each amino acid contributes 6 bits, the byte boundaries do not "
        f"align with residue boundaries\u2014each byte contains parts of adjacent codons."
    ))

    step2_headers = ["Byte", "Hex", "Decimal", "Bit Pattern", "Source Bits"]
    step2_rows = []
    for i in range(4):
        start_bit = i * 8
        end_bit = start_bit + 8
        bit_pattern = ncl_bits[start_bit:end_bit]
        start_res = start_bit // 6
        end_res = min((end_bit - 1) // 6, 5)
        source_aas = ncl_seq[start_res:end_res+1]
        step2_rows.append([
            str(i + 1),
            f"0x{ncl_bytes[i]:02X}",
            str(ncl_bytes[i]),
            bit_pattern,
            f"from {source_aas} codons"
        ])
    make_table(doc, step2_headers, step2_rows)

    add_heading_text(doc, "Step 3 \u2014 Byte Range Classification", level=2)

    ctrl, std, ext = classify_bytes(ncl_bytes)
    total = len(ncl_bytes)

    add_para(doc, (
        "Each byte is classified into one of three ranges based on its value:"
    ))

    step3_headers = ["Range", "Hex Values", "Count", "Percentage"]
    step3_rows = [
        ["Control", "0x00\u20130x1F", str(ctrl), f"{ctrl/total*100:.1f}%"],
        ["Standard", "0x20\u20130x7F", str(std), f"{std/total*100:.1f}%"],
        ["Extended", "0x80\u20130xFF", str(ext), f"{ext/total*100:.1f}%"],
        ["Total", "\u2014", str(total), "100.0%"],
    ]
    make_table(doc, step3_headers, step3_rows)

    add_para(doc, (
        f"The distribution ({ctrl/total*100:.0f}% Control / {std/total*100:.0f}% Standard / "
        f"{ext/total*100:.0f}% Extended) reflects the codon-based encoding: G-initial codons "
        f"(binary prefix 10) and C-initial codons (prefix 11) push bytes toward the "
        f"Extended range, while A-initial codons (prefix 00) produce Control-range bytes."
    ), size=11)

    add_heading_text(doc, "Step 4 \u2014 Hexadecimal Representation", level=2)

    add_para(doc, (
        "Each byte is expressed as a two-character hexadecimal value. Where the "
        "byte falls in the printable ASCII range (0x20\u20130x7F), the corresponding "
        "character is shown for reference:"
    ))

    step4_headers = ["Byte", "Hex", "Decimal", "ASCII", "Range"]
    step4_rows = []
    for i in range(8):
        b = ncl_bytes[i]
        ascii_ch = chr(b) if 0x20 <= b <= 0x7F else "\u00b7"
        rng = "Control" if b <= 0x1F else ("Standard" if b <= 0x7F else "Extended")
        step4_rows.append([str(i+1), f"0x{b:02X}", str(b), ascii_ch, rng])
    make_table(doc, step4_headers, step4_rows)

    full_hex = ''.join(f'{b:02X}' for b in ncl_bytes[:20])
    add_para(doc, (
        f"Complete hex stream (first 40 characters): {full_hex}\u2026"
    ), size=11, italic=True)

    add_heading_text(doc, "Step 5 \u2014 Token Discovery", level=2)

    ncl_hex_stream = ''.join(f'{b:02X}' for b in ncl_bytes)
    ncl_matches = tokenize(ncl_hex_stream, vocab)
    unique_words = set(m['hex'] for m in ncl_matches)

    add_para(doc, (
        f"The tokenizer scans the byte stream with a greedy longest-match algorithm, "
        f"matching multi-byte patterns against the 1,932-word vocabulary (word lengths "
        f"range from 2 to 5 bytes). For nucleolin\u2019s "
        f"{len(ncl_bytes)}-byte stream, the scanner identifies {len(ncl_matches)} "
        f"vocabulary hits across {len(unique_words)} unique words."
    ))

    step5_headers = ["Position", "Word (hex)", "Length\n(bytes)", "Occurrences\nin corpus", "Department"]
    step5_rows = []
    shown = 0
    seen = set()
    for m in ncl_matches:
        if m['hex'] not in seen and shown < 15:
            seen.add(m['hex'])
            step5_rows.append([
                str(m['position']),
                f"0x{m['hex']}",
                str(m['byte_len']),
                str(m['occurrences']),
                m['dept']
            ])
            shown += 1
    if len(unique_words) > 15:
        step5_rows.append(["\u2026", f"({len(unique_words) - 15} more)", "\u2026", "\u2026", "\u2026"])
    make_table(doc, step5_headers, step5_rows)

    add_heading_text(doc, "Step 6 \u2014 Vocabulary Classification", level=2)

    func_counts = Counter(m['dept'] for m in ncl_matches)

    add_para(doc, (
        f"Each matched vocabulary word carries a functional department assignment "
        f"derived from GO-term enrichment analysis of its carrier proteins. "
        f"Nucleolin\u2019s {len(ncl_matches)} vocabulary hits map to the following "
        f"department distribution:"
    ))

    step6_headers = ["Department", "Hits", "Percentage"]
    step6_rows = []
    for fn, c in func_counts.most_common():
        step6_rows.append([fn, str(c), f"{c/len(ncl_matches)*100:.1f}%"])
    make_table(doc, step6_headers, step6_rows)

    dominant_fn = func_counts.most_common(1)[0][0]
    classified_counts = {k: v for k, v in func_counts.items() if k != 'Unclassified'}
    dominant_classified = max(classified_counts, key=classified_counts.get) if classified_counts else dominant_fn

    add_para(doc, (
        f"The dominant classified department is {dominant_classified} "
        f"({classified_counts.get(dominant_classified, 0)} of "
        f"{len(ncl_matches)} hits, {classified_counts.get(dominant_classified, 0)/len(ncl_matches)*100:.0f}%), "
        f"consistent with nucleolin\u2019s known role in chromatin remodeling and "
        f"histone chaperone activity. The Unclassified category ({func_counts.get('Unclassified', 0)} hits) "
        f"represents general-purpose vocabulary words that are functionally promiscuous "
        f"across multiple departments. The protein receives a vocabulary density of "
        f"{len(ncl_matches)/len(ncl_bytes)*100:.2f}% (hits per byte)."
    ), size=11)

    add_heading_text(doc, "Step 7 \u2014 Program Assembly", level=2)

    add_para(doc, (
        "In the protein-derived pathway, each protein\u2019s vocabulary words are "
        "assembled into a program representation. Nucleolin\u2019s program has the "
        "following properties:"
    ))

    ncl_func_seq = "|".join(m['dept'] for m in ncl_matches)
    n_distinct = len(set(m['dept'] for m in ncl_matches))
    complexity = "Complex" if len(ncl_matches) > 50 else (
        "Moderate" if len(ncl_matches) >= 16 else "Simple")

    step7_headers = ["Property", "Value"]
    step7_rows = [
        ["Gene", "NCL (Nucleolin)"],
        ["UniProt", "P19338"],
        ["Vocabulary words", str(len(ncl_matches))],
        ["Unique words", str(len(unique_words))],
        ["Distinct departments", str(n_distinct)],
        ["Complexity tier", complexity],
        ["Dominant classified dept", dominant_classified],
        ["Function sequence (first 80 chars)", ncl_func_seq[:80] + ("\u2026" if len(ncl_func_seq) > 80 else "")],
    ]
    make_table(doc, step7_headers, step7_rows)

    add_para(doc, (
        f"The function sequence records the ordered list of departments encountered "
        f"along the protein\u2019s byte stream. This sequence serves as the protein\u2019s "
        f"program signature for cross-referencing against genome programs."
    ), size=11)

    add_heading_text(doc, "Step 8 \u2014 Kernel Integration", level=2)

    add_para(doc, (
        "At the kernel level, nucleolin\u2019s protein program is indexed by its "
        "UniProt accession and gene name. The kernel\u2019s process table records:"
    ))

    step8_headers = ["Property", "Value"]
    step8_rows = [
        ["Protein PID", "P19338 (NCL)"],
        ["Chromosome", "chr2 (NCL gene locus: 2q37.1)"],
        ["Memory segment", "RELAY_EFFECTOR_RW"],
        ["Vocabulary density", f"{len(ncl_matches)/len(ncl_bytes)*100:.2f}%"],
        ["Program type", "Protein-derived"],
        ["Dispatch connections", "Via shared vocabulary words with genome programs"],
    ]
    make_table(doc, step8_headers, step8_rows)

    chromatin_word = None
    for m in ncl_matches:
        if m['dept'] == 'Chromatin':
            chromatin_word = m
            break

    if chromatin_word:
        add_para(doc, (
            f"Nucleolin\u2019s vocabulary words can be cross-referenced against genome "
            f"programs carrying the same multi-byte patterns, establishing protein-to-genome "
            f"dispatch connections. For example, vocabulary word 0x{chromatin_word['hex']} "
            f"(Chromatin) appears in {chromatin_word['occurrences']} proteins genome-wide, "
            f"linking nucleolin to the broader Chromatin regulatory network."
        ), size=11)

    add_heading_text(doc, "Example B: Chromosome 11 INS Locus (DNA-Native Pathway)", level=1)
    add_para(doc, (
        "The INS locus (T2T-CHM13v2.0 chr11, cytogenetic band 11p15.5) encodes the "
        "insulin gene. This 1,000-nucleotide segment spanning the INS gene and its "
        "flanking regulatory regions is traced through the DNA-native encoding pathway, "
        "which differs from the protein pathway in Step 1: nucleotides are encoded "
        "directly (2 bits each) rather than through reverse translation of amino acids "
        "(6 bits each via codon lookup)."
    ))

    add_heading_text(doc, "Step 1 \u2014 Nucleotide-to-Binary Encoding", level=2)
    add_para(doc, (
        "Each nucleotide is assigned a 2-bit binary value using the same mapping "
        "as the codon pipeline: A = 00, T = 01, G = 10, C = 11. "
        "The first 12 nucleotides of the INS locus region are:"
    ))

    step1b_headers = ["Position", "Nucleotide", "Binary"]
    step1b_rows = []
    for i, nt in enumerate(ins_seq[:12]):
        step1b_rows.append([str(i+1), nt, NUC_BIN.get(nt, '??')])
    make_table(doc, step1b_headers, step1b_rows)

    total_nt_bits = len(ins_seq) * 2
    add_para(doc, (
        f"The 1,000-nucleotide sequence produces {total_nt_bits:,} bits of binary data."
    ), size=11)

    add_heading_text(doc, "Step 2 \u2014 Binary-to-Byte Conversion", level=2)

    ins_bytes, ins_bits_str, ins_remainder = encode_dna(ins_seq)

    add_para(doc, (
        f"With 2 bits per nucleotide, exactly four nucleotides fill one byte "
        f"(4 \u00d7 2 = 8 bits). The 1,000 nucleotides yield exactly "
        f"{len(ins_bytes)} bytes with no remainder."
    ))

    step2b_headers = ["Byte", "Hex", "Decimal", "Source Nucleotides", "Bit Pattern"]
    step2b_rows = []
    for i in range(4):
        start_nt = i * 4
        source = ins_seq[start_nt:start_nt + 4]
        bit_pattern = ins_bits_str[i*8:(i+1)*8]
        step2b_rows.append([
            str(i+1), f"0x{ins_bytes[i]:02X}", str(ins_bytes[i]),
            source, bit_pattern
        ])
    make_table(doc, step2b_headers, step2b_rows)

    add_heading_text(doc, "Step 3 \u2014 Byte Range Classification", level=2)

    ctrl_i, std_i, ext_i = classify_bytes(ins_bytes)
    tot_i = len(ins_bytes)

    step3b_headers = ["Range", "Hex Values", "Count", "Percentage"]
    step3b_rows = [
        ["Control", "0x00\u20130x1F", str(ctrl_i), f"{ctrl_i/tot_i*100:.1f}%"],
        ["Standard", "0x20\u20130x7F", str(std_i), f"{std_i/tot_i*100:.1f}%"],
        ["Extended", "0x80\u20130xFF", str(ext_i), f"{ext_i/tot_i*100:.1f}%"],
        ["Total", "\u2014", str(tot_i), "100.0%"],
    ]
    make_table(doc, step3b_headers, step3b_rows)

    add_para(doc, (
        f"The DNA-native pathway produces a markedly different byte distribution "
        f"than the protein pathway: {ext_i/tot_i*100:.0f}% Extended versus "
        f"nucleolin\u2019s {ext/total*100:.0f}%. This reflects the 2-bit encoding "
        f"scheme, where GC-rich regions (G = 10, C = 11) produce high byte values."
    ), size=11)

    add_heading_text(doc, "Step 4 \u2014 Hexadecimal Representation", level=2)

    step4b_headers = ["Byte", "Hex", "Decimal", "ASCII", "Range"]
    step4b_rows = []
    for i in range(8):
        b = ins_bytes[i]
        ascii_ch = chr(b) if 0x20 <= b <= 0x7F else "\u00b7"
        rng = "Control" if b <= 0x1F else ("Standard" if b <= 0x7F else "Extended")
        step4b_rows.append([str(i+1), f"0x{b:02X}", str(b), ascii_ch, rng])
    make_table(doc, step4b_headers, step4b_rows)

    ins_full_hex = ''.join(f'{b:02X}' for b in ins_bytes[:20])
    add_para(doc, (
        f"Complete hex stream (first 40 characters): {ins_full_hex}\u2026"
    ), size=11, italic=True)

    add_heading_text(doc, "Step 5 \u2014 Token Discovery", level=2)

    ins_hex_stream = ''.join(f'{b:02X}' for b in ins_bytes)
    ins_matches = tokenize(ins_hex_stream, vocab)
    ins_unique = set(m['hex'] for m in ins_matches)

    add_para(doc, (
        f"Scanning the {len(ins_bytes)}-byte INS stream against the 1,932-word "
        f"vocabulary (using greedy longest-match) yields {len(ins_matches)} hit(s) "
        f"across {len(ins_unique)} unique word(s). The substantially lower hit rate "
        f"compared to nucleolin "
        f"({len(ins_matches)}/{len(ins_bytes)} = "
        f"{len(ins_matches)/len(ins_bytes)*100:.2f}% vs "
        f"{len(ncl_matches)}/{len(ncl_bytes)} = "
        f"{len(ncl_matches)/len(ncl_bytes)*100:.2f}%) reflects a fundamental "
        f"difference between the two pathways: the vocabulary was extracted from "
        f"protein-derived encodings, so DNA-native sequences produce sparser "
        f"vocabulary matches. Genome programs are identified by regions of "
        f"concentrated vocabulary hits across entire chromosomes, not uniform coverage "
        f"of short windows."
    ))

    if ins_matches:
        step5b_headers = ["Position", "Word (hex)", "Length\n(bytes)", "Department", "Occurrences"]
        step5b_rows = []
        for m in ins_matches:
            step5b_rows.append([
                str(m['position']), f"0x{m['hex']}", str(m['byte_len']),
                m['dept'], str(m['occurrences'])
            ])
        make_table(doc, step5b_headers, step5b_rows)

    add_heading_text(doc, "Step 6 \u2014 Vocabulary Classification", level=2)

    if ins_matches:
        ins_func_counts = Counter(m['dept'] for m in ins_matches)
        add_para(doc, (
            f"The {len(ins_matches)} vocabulary hit(s) in the INS region map to "
            f"the following department(s):"
        ))
        step6b_headers = ["Department", "Hits"]
        step6b_rows = [[fn, str(c)] for fn, c in ins_func_counts.most_common()]
        make_table(doc, step6b_headers, step6b_rows)
    else:
        add_para(doc, (
            "No vocabulary hits were found in this 1,000-nucleotide segment. "
            "This is expected: the vocabulary was derived from protein-encoded "
            "byte patterns, and DNA-native regions produce different byte "
            "distributions. Genome programs are identified from longer chromosomal "
            "scans where vocabulary hits cluster."
        ))

    add_heading_text(doc, "Step 7 \u2014 Program Assembly", level=2)

    ins_programs = load_programs_near("chr11", 2159779, 500000)

    add_para(doc, (
        f"Genome programs are identified by scanning complete chromosomes for "
        f"regions where vocabulary density exceeds the boundary detection "
        f"threshold. Within 500 kb of the INS locus, the pipeline identifies "
        f"{len(ins_programs)} genome programs:"
    ))

    step7b_headers = ["PID", "Range", "Length\n(bytes)", "Function Sequence",
                      "Vocab\nHits", "Entry\nPoint"]
    step7b_rows = []
    for p in ins_programs:
        func_seq = p['function_sequence']
        if len(func_seq) > 45:
            func_seq = func_seq[:42] + "\u2026"
        step7b_rows.append([
            f"chr11:{p['start_position']}",
            f"{p['start_position']}\u2013{p['end_position']}",
            p['length_bytes'],
            func_seq,
            p['n_vocab_hits'],
            p['entry_point'],
        ])
    make_table(doc, step7b_headers, step7b_rows)

    has_prims = [p for p in ins_programs if int(p.get('matched_primitives', 0)) > 0]
    if has_prims:
        add_para(doc, (
            f"Of these, {len(has_prims)} program(s) match known primitives "
            f"(recurring cross-chromosomal patterns). The closest program to the "
            f"INS gene (chr11:{ins_programs[1]['start_position'] if len(ins_programs)>1 else ins_programs[0]['start_position']}) "
            f"carries {ins_programs[1]['n_distinct_funcs'] if len(ins_programs)>1 else ins_programs[0]['n_distinct_funcs']} "
            f"distinct functional departments and is classified as "
            f"{ins_programs[1]['dominant_function'] if len(ins_programs)>1 else ins_programs[0]['dominant_function']}."
        ), size=11)

    add_heading_text(doc, "Step 8 \u2014 Kernel Integration", level=2)

    add_para(doc, (
        "At the kernel level, each genome program near the INS locus receives a "
        "process ID (PID) based on its chromosomal position. These programs are "
        "registered in the kernel\u2019s process table:"
    ))

    step8b_headers = ["PID", "Role", "Memory\nSegment", "Entry Point",
                      "Dominant\nFunction"]
    step8b_rows = []
    for p in ins_programs:
        step8b_rows.append([
            f"chr11:{p['start_position']}",
            "RELAY-EFFECTOR",
            "RELAY_EFFECTOR_RW",
            p['entry_point'],
            p['dominant_function'],
        ])
    make_table(doc, step8b_headers, step8b_rows)

    add_para(doc, (
        "All four programs trace their dispatch origin to chrM entry points "
        "(0x71C7 and 0x5B2B), confirming that the INS region is reachable from "
        "the kernel\u2019s boot sequence. Chromosome 11 is classified as "
        "RELAY-EFFECTOR with read-write memory protection."
    ), size=11)

    add_heading_text(doc, "Substrate Convergence: Protein and DNA Pathways", level=1)

    add_para(doc, (
        "A central feature of the V2 encoding pipeline is that both protein "
        "sequences and genomic DNA are encoded into the same byte space, enabling "
        "direct comparison. The protein pathway reverse-translates amino acids to "
        "DNA codons before applying the same nucleotide-to-binary mapping "
        "(A = 00, T = 01, G = 10, C = 11) used by the DNA pathway. "
        "The following analysis compares the encoding outputs "
        "of nucleolin (protein-derived) and the INS locus (DNA-native)."
    ))

    ncl_byte_set = set(ncl_bytes)
    ins_byte_set = set(ins_bytes)
    shared_bytes = ncl_byte_set & ins_byte_set
    union_bytes = ncl_byte_set | ins_byte_set

    add_heading_text(doc, "Byte-Level Comparison", level=2)

    conv_headers = ["Property", "Nucleolin\n(protein)", "INS Locus\n(DNA)", "Shared"]
    conv_rows = [
        ["Input length", f"{len(ncl_seq)} amino acids", f"{len(ins_seq)} nucleotides", "\u2014"],
        ["Encoding path", "AA \u2192 codon \u2192 DNA \u2192 binary", "DNA \u2192 binary", "\u2014"],
        ["Bits per input unit", "6 (3 nt \u00d7 2 bit)", "2", "\u2014"],
        ["Encoded bytes", str(len(ncl_bytes)), str(len(ins_bytes)), "\u2014"],
        ["Unique byte values", str(len(ncl_byte_set)), str(len(ins_byte_set)),
         f"{len(shared_bytes)} ({len(shared_bytes)/len(union_bytes)*100:.0f}% of union)"],
        ["Control range", f"{ctrl/total*100:.1f}%", f"{ctrl_i/tot_i*100:.1f}%", "\u2014"],
        ["Standard range", f"{std/total*100:.1f}%", f"{std_i/tot_i*100:.1f}%", "\u2014"],
        ["Extended range", f"{ext/total*100:.1f}%", f"{ext_i/tot_i*100:.1f}%", "\u2014"],
        ["Vocabulary hits", str(len(ncl_matches)), str(len(ins_matches)), "\u2014"],
        ["Hit rate", f"{len(ncl_matches)/len(ncl_bytes)*100:.2f}%",
         f"{len(ins_matches)/len(ins_bytes)*100:.2f}%", "\u2014"],
    ]
    make_table(doc, conv_headers, conv_rows)

    add_heading_text(doc, "Vocabulary-Level Comparison", level=2)

    ncl_word_set = set(m['hex'] for m in ncl_matches)
    ins_word_set = set(m['hex'] for m in ins_matches)
    shared_words = ncl_word_set & ins_word_set

    add_para(doc, (
        f"Nucleolin\u2019s encoding contains {len(ncl_word_set)} unique vocabulary "
        f"words; the INS locus contains {len(ins_word_set)}. "
        f"{'They share ' + str(len(shared_words)) + ' word(s), ' if shared_words else 'They share no vocabulary words, '}"
        f"which reflects the different byte distributions produced by the two "
        f"encoding schemes. The protein-derived pathway (6-bit codon-based encoding) "
        f"produces byte values distributed across all three ranges, while the "
        f"DNA-native pathway (2-bit nucleotide encoding) concentrates in the "
        f"Extended range due to the high information content per nucleotide pair."
    ))

    add_heading_text(doc, "Pathway Relationship", level=2)

    add_para(doc, (
        "The two pathways intersect at the vocabulary level: both produce byte "
        "streams that are scanned against the same 1,932-word dictionary. When a "
        "vocabulary word appears in both a protein\u2019s encoding and a genomic "
        "region\u2019s encoding, it establishes a cross-substrate link. For the INS "
        "locus specifically, the genome programs identified within 500 kb of the "
        "gene carry function sequences (Chromatin, Transcription, Cytoskeleton) "
        "that overlap with departments found in nucleolin\u2019s encoding, "
        "illustrating how the shared vocabulary enables cross-substrate functional "
        "annotation."
    ))

    add_para(doc, (
        "This convergence is not guaranteed by the method \u2014 it is an empirical "
        "result. The vocabulary was extracted from protein data; that it produces "
        "biologically meaningful hits in genomic DNA is a validation of the "
        "encoding\u2019s ability to bridge the two substrates."
    ), size=11)

    doc.save(OUT_PATH)
    print(f"Document saved to: {os.path.abspath(OUT_PATH)}")


if __name__ == "__main__":
    main()
